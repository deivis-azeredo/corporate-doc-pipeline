import os
import time
from docx import Document
from groq import Groq

# ========================
# CONFIGURAÇÕES E PASTAS
# ========================
GROQ_API_KEY = ""
INPUT_DIR = "input_docs"
OUTPUT_DIR = "output_docs"

# Garante que as pastas de saída existam
os.makedirs(OUTPUT_DIR, exist_ok=True)

print("🔌 Conectando à API da Groq...")
groq_client = Groq(api_key=GROQ_API_KEY)

# ========================
# PROCESSAMENTO EM LOTE (BATCH PROCESSING)
# ========================
print(f"📂 Lendo documentos na pasta '{INPUT_DIR}'...")

# Varre todos os arquivos dentro da pasta de entrada
for filename in os.listdir(INPUT_DIR):
    if filename.endswith(".docx") and not filename.startswith("~"):
        input_path = os.path.join(INPUT_DIR, filename)
        output_path = os.path.join(OUTPUT_DIR, f"revisado_{filename}")
        
        print(f"\n📄 Processando arquivo: {filename}")
        
        doc = Document(input_path)
        new_doc = Document()
        total_paragraphs = len(doc.paragraphs)
        
        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                new_doc.add_paragraph("") 
                continue
                
            print(f"  [{index + 1}/{total_paragraphs}] Analisando parágrafo...")
            
            try:
                response = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are an expert corporate editor. "
                                "Review and polish the following text. Correct grammar, "
                                "typos, and structural issues. Enhance clarity and corporate tone. "
                                "Return ONLY the refined text without any quotes or explanations."
                            )
                        },
                        {
                            "role": "user",
                            "content": text
                        }
                    ]
                )
                
                refined_text = response.choices[0].message.content.strip()
                new_doc.add_paragraph(refined_text)
                time.sleep(2) # Proteção de limite da API
                
            except Exception as e:
                print(f"  ❌ Erro no parágrafo {index + 1}: {e}")
                new_doc.add_paragraph(text) # Salva original em caso de erro

        # Salva o arquivo finalizado direto na pasta de saída
        new_doc.save(output_path)
        print(f"✅ Arquivo salvo com sucesso em: {output_path}")

print("\n🎉 Processamento em lote finalizado!")